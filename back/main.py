from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import pandas as pd
import io
import numpy as np
import math
import os
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import tempfile

app = FastAPI(title="API для статистического анализа координат")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"message": "API для анализа координат работает!"}

def calculate_statistics(df: pd.DataFrame) -> Dict[str, Any]:
    stats = {
        "mean": {
            "x": float(df['x'].mean()),
            "y": float(df['y'].mean()),
            "z": float(df['z'].mean())
        },
        "median": {
            "x": float(df['x'].median()),
            "y": float(df['y'].median()),
            "z": float(df['z'].median())
        },
        "std": {
            "x": float(df['x'].std()),
            "y": float(df['y'].std()),
            "z": float(df['z'].std())
        },
        "min": {
            "x": float(df['x'].min()),
            "y": float(df['y'].min()),
            "z": float(df['z'].min())
        },
        "max": {
            "x": float(df['x'].max()),
            "y": float(df['y'].max()),
            "z": float(df['z'].max())
        }
    }
    return stats

def generate_plots(df: pd.DataFrame) -> str:
    """Генерирует графики и возвращает временный путь к файлу"""
    plt.figure(figsize=(12, 8))
    
    # Гистограмма распределения координат
    plt.subplot(2, 2, 1)
    df[['x', 'y', 'z']].plot.hist(alpha=0.5, bins=20)
    plt.title('Распределение координат')
    
    # Boxplot
    plt.subplot(2, 2, 2)
    df[['x', 'y', 'z']].plot.box()
    plt.title('Boxplot координат')
    
    # 3D scatter plot
    ax = plt.subplot(2, 2, 3, projection='3d')
    ax.scatter(df['x'], df['y'], df['z'])
    ax.set_xlabel('X')
    ax.set_ylabel('Y')
    ax.set_zlabel('Z')
    plt.title('3D визуализация координат')
    
    plt.tight_layout()
    
    temp_dir = tempfile.mkdtemp()
    plot_path = os.path.join(temp_dir, "plots.png")
    plt.savefig(plot_path)
    plt.close()
    
    return plot_path

def generate_word_report(df: pd.DataFrame, stats: Dict[str, Any], plot_path: str) -> io.BytesIO:
    doc = Document()

    # Устанавливаем стандартные стили для документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = 12
    
    # Заголовок
    title = doc.add_heading('Статистический отчет по координатам', level=0)
    title.style = doc.styles['Heading 1']
    
    # Основная информация
    doc.add_paragraph(f"Всего точек: {len(df)}", style='BodyText')
    doc.add_paragraph(f"Дата анализа: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}", style='BodyText')
    
    # Добавляем таблицу с основными статистиками
    doc.add_heading('Основные статистики', level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Метрика'
    hdr_cells[1].text = 'X'
    hdr_cells[2].text = 'Y'
    hdr_cells[3].text = 'Z'
    
    # Заполняем данные
    metrics = ['Среднее', 'Медиана', 'Станд. отклонение', 'Минимум', 'Максимум']
    stat_keys = ['mean', 'median', 'std', 'min', 'max']
    
    for i, (metric, key) in enumerate(zip(metrics, stat_keys)):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = metric
        row_cells[1].text = f"{stats[key]['x']:.2f}"
        row_cells[2].text = f"{stats[key]['y']:.2f}"
        row_cells[3].text = f"{stats[key]['z']:.2f}"
    
    # Добавляем графики
    doc.add_heading('Визуализация данных', level=1)
    doc.add_picture(plot_path, width=Inches(6.0))
    
    # Сохраняем документ в BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

def generate_markdown_report(df: pd.DataFrame, stats: Dict[str, Any]) -> str:
    report = "# Статистический отчет по координатам\n\n"
    report += f"**Всего точек:** {len(df)}\n"
    report += f"**Дата анализа:** {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}\n\n"
    
    report += "## Основные статистики\n\n"
    report += "| Метрика | X | Y | Z |\n"
    report += "|---------|---|---|---|\n"
    report += f"| Среднее | {stats['mean']['x']:.2f} | {stats['mean']['y']:.2f} | {stats['mean']['z']:.2f} |\n"
    report += f"| Медиана | {stats['median']['x']:.2f} | {stats['median']['y']:.2f} | {stats['median']['z']:.2f} |\n"
    report += f"| Станд. отклонение | {stats['std']['x']:.2f} | {stats['std']['y']:.2f} | {stats['std']['z']:.2f} |\n"
    report += f"| Минимум | {stats['min']['x']:.2f} | {stats['min']['y']:.2f} | {stats['min']['z']:.2f} |\n"
    report += f"| Максимум | {stats['max']['x']:.2f} | {stats['max']['y']:.2f} | {stats['max']['z']:.2f} |\n\n"
    
    report += "## Пример данных (первые 5 строк)\n\n"
    report += df.head().to_markdown(index=False) + "\n\n"
    
    report += "## Визуализация\n\n"
    report += "![Графики координат](plots.png)\n"
    
    return report

@app.post("/analyze")
async def analyze_file(file: UploadFile = File(...)) -> Dict[str, Any]:
    try:
        contents = await file.read()
        df = pd.read_excel(io.BytesIO(contents))

        required_columns = ['x', 'y', 'z']
        if not all(col in df.columns for col in required_columns):
            raise HTTPException(status_code=400, detail="Файл должен содержать столбцы x, y, z")
        
        df[['x', 'y', 'z']] = df[['x', 'y', 'z']].astype(float)
        
        # Рассчитываем статистику
        stats = calculate_statistics(df)
        
        # Генерируем графики
        plot_path = generate_plots(df)
        
        # Создаем отчеты
        word_report = generate_word_report(df, stats, plot_path)
        markdown_report = generate_markdown_report(df, stats)
        
        # Читаем графики как base64 для фронтенда
        with open(plot_path, "rb") as image_file:
            encoded_plots = base64.b64encode(image_file.read()).decode('utf-8')
        
        # Удаляем временные файлы
        os.remove(plot_path)
        os.rmdir(os.path.dirname(plot_path))
        
        return {
            "status": "success",
            "statistics": stats,
            "markdown_report": markdown_report,
            "plots_base64": encoded_plots,
            "word_report": base64.b64encode(word_report.read()).decode('utf-8')
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)